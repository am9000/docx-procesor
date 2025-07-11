from docx import Document

def process_docx(file_path):
    """Process a .docx file and extract its text content."""
    if not file_path.endswith('.docx'):
        raise ValueError("The provided file is not a .docx file.")
        
    try:
        # Load the document
        doc = Document(file_path)
        
        # Extract text from the document
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        
        return '\n'.join(text)
    
    except Exception as e:
        print(f"An error occurred while processing the document: {e}")
        return None

def add_generated_lines_to_headers(file_path):
    """
    Reads a .docx file, adds generated lines under each h1, h2, h3 header, and saves output as 'output.docx'.
    """
    if not file_path.endswith('.docx'):
        raise ValueError("The provided file is not a .docx file.")
    try:
        doc = Document(file_path)
        tables = doc.tables
        if tables:
            tables_count = len(tables)
            print(f"!!! The document contains {tables_count} table(s) !!!")
            
        header_count = 1
        for para in doc.paragraphs:
            # Check for header (h1, h2, h3)
            style_name = str(para.style.name).lower()
            if style_name in ["heading 1", "heading 2", "heading 3"]:
                para.insert_paragraph_before(f"|ID: SHARED-{header_count}|\n|LOKALIZACJA: SHARED|")
                header_count += 1
        doc.save("output.docx")
        return True
    except Exception as e:
        print(f"An error occurred while processing the document: {e}")
        return False
    
if __name__ == "__main__":
    file_path = 'test-data/SHARED-test-short-input.docx'
    if add_generated_lines_to_headers(file_path):
        print("Document processed and saved as 'output.docx'.")
    else:
        print("Failed to process the document.")